import docx, sys
from lxml import etree
sys.stdout.reconfigure(encoding='utf-8')
doc = docx.Document(r'C:\Users\DinhViet\Documents\DoAnTotNghiep\DATN_CNS_TranDinhViet_FinalVersion_1.2025.docx')

# Check default font in document
print('=== DOCUMENT DEFAULT FONT ===')
ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
rPrDefault = doc.element.find(f'.//{ns}rPrDefault')
if rPrDefault is not None:
    print(etree.tostring(rPrDefault, pretty_print=True).decode())
pPrDefault = doc.element.find(f'.//{ns}pPrDefault')
if pPrDefault is not None:
    print(etree.tostring(pPrDefault, pretty_print=True).decode())

# Check headers/footers for page numbering
print()
print('=== HEADERS AND FOOTERS ===')
for i, sec in enumerate(doc.sections):
    print(f'Section {i}:')
    if sec.header and sec.header.paragraphs:
        for p in sec.header.paragraphs:
            if p.text.strip():
                print(f'  Header: {p.text[:80]}')
    if sec.footer and sec.footer.paragraphs:
        for p in sec.footer.paragraphs:
            txt = p.text.strip()
            print(f'  Footer: "{txt}" | alignment={p.alignment}')
            fields = p._element.findall(f'.//{ns}fldChar')
            if fields:
                print(f'    Has field chars (page numbers): {len(fields)}')
