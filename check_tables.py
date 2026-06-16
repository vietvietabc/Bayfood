import docx, sys
from docx.shared import Pt, Cm, Emu
from docx.oxml.ns import qn
from lxml import etree
sys.stdout.reconfigure(encoding='utf-8')

doc = docx.Document(r'C:\Users\DinhViet\Documents\DoAnTotNghiep\DATN_CNS_TranDinhViet_FinalVersion_Formatted.docx')

# Check tables in detail
print('=== TABLES DETAIL ===')
for t_idx, table in enumerate(doc.tables):
    print(f'\nTable {t_idx}:')
    # Check table style
    print(f'  Style: {table.style.name if table.style else "None"}')
    print(f'  Alignment: {table.alignment}')
    
    # Check borders via XML
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is not None:
        tblBorders = tblPr.find(qn('w:tblBorders'))
        if tblBorders is not None:
            print(f'  Has borders: YES')
        else:
            print(f'  Has borders: NO (in tblPr)')
    
    # Check first row content
    if table.rows:
        first_row = table.rows[0]
        cells_text = []
        for cell in first_row.cells:
            cells_text.append(cell.text[:40])
        print(f'  First row ({len(first_row.cells)} cells): {cells_text}')
        
        # Check font in first cell
        if first_row.cells:
            for p in first_row.cells[0].paragraphs[:1]:
                for run in p.runs[:1]:
                    print(f'  Cell font: name={run.font.name}, size={run.font.size}, bold={run.font.bold}')
                print(f'  Cell alignment: {p.alignment}')
                pf = p.paragraph_format
                print(f'  Cell line_spacing: {pf.line_spacing}')
                print(f'  Cell first_indent: {pf.first_line_indent}')
    
    print(f'  Total rows: {len(table.rows)}')
    if t_idx >= 5:  # Show first 6 tables
        print('  ... (showing first 6 tables only)')
        break

# Check image paragraphs
print('\n\n=== IMAGE PARAGRAPHS ===')
for i, p in enumerate(doc.paragraphs):
    # Check if paragraph has images (inline shapes)
    has_image = False
    drawing_elements = p._element.findall('.//' + qn('w:drawing'))
    if drawing_elements:
        has_image = True
    
    # Also check for old-style images
    pict_elements = p._element.findall('.//' + qn('w:pict'))
    if pict_elements:
        has_image = True
    
    if has_image:
        style = p.style.name if p.style else 'None'
        pf = p.paragraph_format
        print(f'[{i}] Style:{style} | Align:{pf.alignment} | Indent:{pf.first_line_indent} | Text:{p.text[:60]}')
        # Check image size
        for drawing in drawing_elements:
            extent = drawing.find('.//' + qn('wp:extent'))
            if extent is not None:
                cx = int(extent.get('cx', 0))
                cy = int(extent.get('cy', 0))
                print(f'  Image size: {cx/914400:.2f}" x {cy/914400:.2f}" ({cx/360000:.1f}cm x {cy/360000:.1f}cm)')

# Check paragraphs with "Hình" or "Bảng" caption
print('\n\n=== CAPTIONS (Hình/Bảng) ===')
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text and (text.startswith('Hình ') or text.startswith('Bảng ')):
        style = p.style.name if p.style else 'None'
        pf = p.paragraph_format
        font_info = ''
        for run in p.runs[:1]:
            font_info = f'font={run.font.name}, size={run.font.size}, bold={run.font.bold}, italic={run.font.italic}'
        print(f'[{i}] Style:{style} | Align:{pf.alignment} | {font_info} | {text[:80]}')
