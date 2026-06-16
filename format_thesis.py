"""
Script fix format bảng + hình ảnh cho ĐATN
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

import docx
from docx import Document
from docx.shared import Pt, Cm, Inches, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree

INPUT_FILE = r'C:\Users\DinhViet\Documents\DoAnTotNghiep\DATN_CNS_TranDinhViet_FinalVersion_1.2025.docx'
OUTPUT_FILE = r'C:\Users\DinhViet\Documents\DoAnTotNghiep\DATN_CNS_TranDinhViet_FinalVersion_Formatted.docx'

FONT_NAME = 'Times New Roman'
BODY_SIZE = Pt(13)
HEADING1_SIZE = Pt(14)
HEADING2_SIZE = Pt(13)
HEADING3_SIZE = Pt(13)
LINE_SPACING = 1.5
FIRST_LINE_INDENT = Cm(1.27)
SPACE_BEFORE = Pt(6)
SPACE_AFTER = Pt(6)

COVER_STYLES = [
    'Bìa - ĐHĐN', 'Bìa ĐATN', 'Bìa - Ngành', 
    'Bìa - Người hướng dẫn', 'Normal - Lề giữa',
    'TTat-NXet-NVu-LNĐ'
]

HEADING_STYLES = {
    'Heading 1': {'size': HEADING1_SIZE, 'bold': True, 'italic': False, 'alignment': WD_ALIGN_PARAGRAPH.CENTER},
    'Heading 2': {'size': HEADING2_SIZE, 'bold': True, 'italic': False, 'alignment': None},
    'Heading 3': {'size': HEADING3_SIZE, 'bold': True, 'italic': True, 'alignment': None},
    'Heading 4': {'size': HEADING3_SIZE, 'bold': True, 'italic': True, 'alignment': None},
    'Heading 5': {'size': HEADING3_SIZE, 'bold': False, 'italic': True, 'alignment': None},
}


def set_font_for_run(run, font_name=FONT_NAME, font_size=None, bold=None, italic=None):
    run.font.name = font_name
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}" w:eastAsia="{font_name}" w:cs="{font_name}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:eastAsia'), font_name)
        rFonts.set(qn('w:cs'), font_name)
    if font_size is not None:
        run.font.size = font_size
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic


def set_style_font(style, font_name=FONT_NAME, font_size=None, bold=None, italic=None):
    style.font.name = font_name
    rPr = style.element.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        style.element.append(rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}" w:eastAsia="{font_name}" w:cs="{font_name}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:eastAsia'), font_name)
        rFonts.set(qn('w:cs'), font_name)
    if font_size is not None:
        style.font.size = font_size
    if bold is not None:
        style.font.bold = bold
    if italic is not None:
        style.font.italic = italic


def set_document_default_font(doc, font_name, font_size):
    styles_element = doc.styles.element
    docDefaults = styles_element.find(qn('w:docDefaults'))
    if docDefaults is None:
        docDefaults = parse_xml(f'<w:docDefaults {nsdecls("w")}></w:docDefaults>')
        styles_element.insert(0, docDefaults)
    rPrDefault = docDefaults.find(qn('w:rPrDefault'))
    if rPrDefault is None:
        rPrDefault = parse_xml(f'<w:rPrDefault {nsdecls("w")}></w:rPrDefault>')
        docDefaults.insert(0, rPrDefault)
    rPr = rPrDefault.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        rPrDefault.append(rPr)
    for child in rPr.findall(qn('w:rFonts')):
        rPr.remove(child)
    for child in rPr.findall(qn('w:sz')):
        rPr.remove(child)
    for child in rPr.findall(qn('w:szCs')):
        rPr.remove(child)
    rFonts = parse_xml(
        f'<w:rFonts {nsdecls("w")} '
        f'w:ascii="{font_name}" w:hAnsi="{font_name}" '
        f'w:eastAsia="{font_name}" w:cs="{font_name}"/>'
    )
    rPr.insert(0, rFonts)
    size_val = str(int(font_size.pt * 2))
    rPr.append(parse_xml(f'<w:sz {nsdecls("w")} w:val="{size_val}"/>'))
    rPr.append(parse_xml(f'<w:szCs {nsdecls("w")} w:val="{size_val}"/>'))


def set_table_borders(table):
    """Set borders for all cells in a table - full grid borders."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
        tbl.insert(0, tblPr)
    
    # Remove existing borders
    existing_borders = tblPr.find(qn('w:tblBorders'))
    if existing_borders is not None:
        tblPr.remove(existing_borders)
    
    # Add full borders (single line, 4pt = 0.5pt, black)
    borders_xml = f'''<w:tblBorders {nsdecls("w")}>
        <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>
        <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>
        <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>
        <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>
        <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>
        <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>
    </w:tblBorders>'''
    tblPr.append(parse_xml(borders_xml))


def set_cell_borders(cell):
    """Set borders for a single cell."""
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}></w:tcPr>')
        tc.insert(0, tcPr)
    
    existing = tcPr.find(qn('w:tcBorders'))
    if existing is not None:
        tcPr.remove(existing)
    
    borders_xml = f'''<w:tcBorders {nsdecls("w")}>
        <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>
        <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>
        <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>
        <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>
    </w:tcBorders>'''
    tcPr.append(parse_xml(borders_xml))


def format_thesis():
    print(f'Đang mở file: {INPUT_FILE}')
    doc = Document(INPUT_FILE)
    
    # =========================================================
    # 1. Document defaults
    # =========================================================
    print('\n[1/9] Font mặc định...')
    set_document_default_font(doc, FONT_NAME, BODY_SIZE)
    print('  ✓ Times New Roman 13pt')
    
    # =========================================================
    # 2. Page margins
    # =========================================================
    print('\n[2/9] Lề trang...')
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
    print(f'  ✓ {len(doc.sections)} sections: T2-D2-L3-R2 cm')
    
    # =========================================================
    # 3. Style definitions
    # =========================================================
    print('\n[3/9] Cập nhật styles...')
    
    # Normal
    ns = doc.styles['Normal']
    set_style_font(ns, FONT_NAME, BODY_SIZE)
    ns.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    ns.paragraph_format.line_spacing = LINE_SPACING
    ns.paragraph_format.space_before = SPACE_BEFORE
    ns.paragraph_format.space_after = SPACE_AFTER
    ns.paragraph_format.first_line_indent = FIRST_LINE_INDENT
    ns.font.color.rgb = RGBColor(0, 0, 0)
    
    # Headings
    for sname, cfg in HEADING_STYLES.items():
        try:
            s = doc.styles[sname]
            set_style_font(s, FONT_NAME, cfg['size'], cfg['bold'], cfg['italic'])
            s.font.color.rgb = RGBColor(0, 0, 0)
            pf = s.paragraph_format
            pf.alignment = cfg['alignment'] if cfg['alignment'] else WD_ALIGN_PARAGRAPH.LEFT
            pf.line_spacing = LINE_SPACING
            pf.space_before = Pt(12)
            pf.space_after = Pt(6)
            pf.first_line_indent = Cm(0)
            # Remove theme colors
            rPr = s.element.find(qn('w:rPr'))
            if rPr is not None:
                for c in rPr.findall(qn('w:color')):
                    c.set(qn('w:val'), '000000')
                    for a in ['w:themeColor', 'w:themeShade', 'w:themeTint']:
                        try: del c.attrib[qn(a)]
                        except: pass
            print(f'  ✓ {sname}')
        except KeyError:
            pass
    
    # modau
    try:
        s = doc.styles['modau']
        set_style_font(s, FONT_NAME, Pt(14), bold=True)
        s.font.color.rgb = RGBColor(0, 0, 0)
        pf = s.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.line_spacing = LINE_SPACING
        pf.space_before = Pt(12)
        pf.space_after = Pt(6)
        pf.first_line_indent = Cm(0)
        print('  ✓ modau')
    except: pass
    
    # Mở đầu
    try:
        s = doc.styles['Mở đầu']
        set_style_font(s, FONT_NAME, Pt(13), bold=True)
        s.font.color.rgb = RGBColor(0, 0, 0)
        pf = s.paragraph_format
        pf.line_spacing = LINE_SPACING
        pf.space_before = Pt(6)
        pf.space_after = Pt(6)
        pf.first_line_indent = Cm(0)
        print('  ✓ Mở đầu')
    except: pass
    
    # Hình ảnh
    try:
        s = doc.styles['Hình ảnh']
        set_style_font(s, FONT_NAME, BODY_SIZE)
        s.font.color.rgb = RGBColor(0, 0, 0)
        pf = s.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.line_spacing = LINE_SPACING
        pf.first_line_indent = Cm(0)
        print('  ✓ Hình ảnh')
    except: pass
    
    # table of figures
    try:
        s = doc.styles['table of figures']
        set_style_font(s, FONT_NAME, BODY_SIZE)
        s.font.color.rgb = RGBColor(0, 0, 0)
        print('  ✓ table of figures')
    except: pass
    
    # Cover styles
    for sname in COVER_STYLES:
        try:
            s = doc.styles[sname]
            old_size = s.font.size
            set_style_font(s, FONT_NAME)
            if old_size:
                s.font.size = old_size
            s.font.color.rgb = RGBColor(0, 0, 0)
            print(f'  ✓ {sname} (size preserved)')
        except: pass
    
    # Normal - size 14
    try:
        s = doc.styles['Normal - size 14']
        set_style_font(s, FONT_NAME, Pt(14))
        s.font.color.rgb = RGBColor(0, 0, 0)
        s.paragraph_format.line_spacing = LINE_SPACING
        print('  ✓ Normal - size 14')
    except: pass
    
    # Nội dung hd3
    try:
        s = doc.styles['Nội dung hd3']
        set_style_font(s, FONT_NAME, BODY_SIZE)
        s.font.color.rgb = RGBColor(0, 0, 0)
        s.paragraph_format.line_spacing = LINE_SPACING
        s.paragraph_format.first_line_indent = FIRST_LINE_INDENT
        print('  ✓ Nội dung hd3')
    except: pass
    
    # =========================================================
    # 4. Format all paragraphs
    # =========================================================
    print('\n[4/9] Format đoạn văn...')
    para_count = 0
    for i, p in enumerate(doc.paragraphs):
        style_name = p.style.name if p.style else 'Normal'
        
        # Check if paragraph has images
        has_image = bool(p._element.findall('.//' + qn('w:drawing'))) or bool(p._element.findall('.//' + qn('w:pict')))
        
        for run in p.runs:
            set_font_for_run(run, font_name=FONT_NAME)
            
            if style_name in COVER_STYLES:
                pass  # keep original size
            elif style_name == 'modau':
                set_font_for_run(run, font_name=FONT_NAME, font_size=Pt(14))
            elif style_name in HEADING_STYLES:
                cfg = HEADING_STYLES[style_name]
                set_font_for_run(run, font_name=FONT_NAME, font_size=cfg['size'])
            elif style_name == 'Normal - size 14':
                set_font_for_run(run, font_name=FONT_NAME, font_size=Pt(14))
            else:
                set_font_for_run(run, font_name=FONT_NAME, font_size=BODY_SIZE)
            
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.highlight_color = None
        
        # Fix image paragraphs: center, no indent
        if has_image:
            pf = p.paragraph_format
            pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.first_line_indent = Cm(0)
        
        para_count += 1
    print(f'  ✓ {para_count} đoạn văn')
    
    # =========================================================
    # 5. Format captions (Hình x.x / Bảng x.x)
    # =========================================================
    print('\n[5/9] Format caption hình/bảng...')
    caption_count = 0
    for p in doc.paragraphs:
        text = p.text.strip()
        if text and (text.startswith('Hình ') or text.startswith('Bảng ')):
            # Check if it's a real caption (has number pattern like "Hình 1.1" or "Bảng 2.3")
            import re
            if re.match(r'^(Hình|Bảng)\s+\d+\.\d+', text):
                pf = p.paragraph_format
                pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pf.first_line_indent = Cm(0)
                pf.space_before = Pt(6)
                pf.space_after = Pt(6)
                
                for run in p.runs:
                    run.font.italic = True
                    run.font.size = BODY_SIZE
                    run.font.name = FONT_NAME
                    set_font_for_run(run, FONT_NAME, BODY_SIZE, italic=True)
                
                caption_count += 1
    print(f'  ✓ {caption_count} captions (Center, Italic, 13pt)')
    
    # =========================================================
    # 6. FORMAT BẢNG (TABLES) - ĐÂY LÀ PHẦN QUAN TRỌNG
    # =========================================================
    print('\n[6/9] Format bảng biểu...')
    table_count = 0
    for t_idx, table in enumerate(doc.tables):
        # --- A) Set table alignment to CENTER ---
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # --- B) Set table-level borders (full grid) ---
        set_table_borders(table)
        
        # --- C) Auto-fit table width ---
        tbl = table._tbl
        tblPr = tbl.find(qn('w:tblPr'))
        
        # Set table width to 100% of page
        tblW = tblPr.find(qn('w:tblW'))
        if tblW is None:
            tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
            tblPr.append(tblW)
        else:
            tblW.set(qn('w:w'), '5000')
            tblW.set(qn('w:type'), 'pct')
        
        # --- D) Format each cell ---
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                # Set cell-level borders
                set_cell_borders(cell)
                
                # Set vertical alignment to center
                tc = cell._tc
                tcPr = tc.find(qn('w:tcPr'))
                if tcPr is None:
                    tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}></w:tcPr>')
                    tc.insert(0, tcPr)
                vAlign = tcPr.find(qn('w:vAlign'))
                if vAlign is None:
                    vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
                    tcPr.append(vAlign)
                else:
                    vAlign.set(qn('w:val'), 'center')
                
                # Set cell margins (padding)
                tcMar = tcPr.find(qn('w:tcMar'))
                if tcMar is None:
                    tcMar_xml = f'''<w:tcMar {nsdecls("w")}>
                        <w:top w:w="28" w:type="dxa"/>
                        <w:left w:w="57" w:type="dxa"/>
                        <w:bottom w:w="28" w:type="dxa"/>
                        <w:right w:w="57" w:type="dxa"/>
                    </w:tcMar>'''
                    tcPr.append(parse_xml(tcMar_xml))
                
                # Format paragraphs in cell
                for p in cell.paragraphs:
                    pf = p.paragraph_format
                    pf.first_line_indent = Cm(0)
                    pf.line_spacing = 1.15
                    pf.space_before = Pt(2)
                    pf.space_after = Pt(2)
                    
                    # Header row: center + bold
                    if row_idx == 0:
                        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in p.runs:
                            set_font_for_run(run, FONT_NAME, Pt(12), bold=True)
                            run.font.color.rgb = RGBColor(0, 0, 0)
                    else:
                        # Body rows
                        for run in p.runs:
                            set_font_for_run(run, FONT_NAME, Pt(12))
                            run.font.color.rgb = RGBColor(0, 0, 0)
        
        table_count += 1
    print(f'  ✓ {table_count} bảng: borders, center, font 12pt, header bold')
    
    # =========================================================
    # 7. Fix headers and footers
    # =========================================================
    print('\n[7/9] Fix Header/Footer...')
    for i, sec in enumerate(doc.sections):
        # Headers
        if sec.header and sec.header.paragraphs:
            for p in sec.header.paragraphs:
                for run in p.runs:
                    set_font_for_run(run, FONT_NAME, Pt(11), italic=True)
        
        # Footers
        if sec.footer and sec.footer.paragraphs:
            for p in sec.footer.paragraphs:
                if 'SVTH:' in p.text and '……' in p.text:
                    for run in p.runs:
                        if '……' in run.text:
                            run.text = run.text.replace('…….', 'Trần Đình Việt')
                            run.text = run.text.replace('……', 'Trần Đình Việt')
                for run in p.runs:
                    set_font_for_run(run, FONT_NAME, Pt(11))
    print('  ✓ Header/Footer: Times New Roman 11pt')
    
    # =========================================================
    # 8. Fix dates
    # =========================================================
    print('\n[8/9] Sửa thời gian...')
    for p in doc.paragraphs:
        for run in p.runs:
            if '2022' in run.text:
                run.text = run.text.replace('2022', '2026')
                print(f'  ✓ 2022 → 2026')
            if 'tháng …' in run.text:
                run.text = run.text.replace('tháng …', 'tháng 6')
                print(f'  ✓ tháng … → tháng 6')
    
    # =========================================================
    # 9. Final cleanup for image paragraphs
    # =========================================================
    print('\n[9/9] Final cleanup hình ảnh...')
    img_count = 0
    for p in doc.paragraphs:
        has_drawing = bool(p._element.findall('.//' + qn('w:drawing')))
        has_pict = bool(p._element.findall('.//' + qn('w:pict')))
        
        if has_drawing or has_pict:
            pf = p.paragraph_format
            pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.first_line_indent = Cm(0)
            pf.space_before = Pt(6)
            pf.space_after = Pt(6)
            img_count += 1
    print(f'  ✓ {img_count} hình ảnh: Center, no indent')
    
    # =========================================================
    # SAVE
    # =========================================================
    print(f'\n{"="*60}')
    doc.save(OUTPUT_FILE)
    print(f'✅ HOÀN TẤT! Đã lưu: {OUTPUT_FILE}')
    print()
    print('📋 Format đã áp dụng:')
    print('   ┌──────────────────────┬────────────────────────────────┐')
    print('   │ Thuộc tính           │ Giá trị                       │')
    print('   ├──────────────────────┼────────────────────────────────┤')
    print('   │ Font                 │ Times New Roman                │')
    print('   │ Cỡ chữ nội dung     │ 13pt                           │')
    print('   │ Cỡ chữ bảng         │ 12pt                           │')
    print('   │ Heading 1            │ 14pt Bold Center               │')
    print('   │ Heading 2            │ 13pt Bold                      │')
    print('   │ Heading 3            │ 13pt Bold Italic               │')
    print('   │ Dãn dòng            │ 1.5                            │')
    print('   │ Thụt đầu dòng       │ 1.27cm                         │')
    print('   │ Lề trang            │ T2-D2-L3-R2 cm                 │')
    print('   │ Căn lề              │ Justify                        │')
    print('   │ Bảng                │ Borders đầy đủ, Center, 12pt   │')
    print('   │ Bảng header         │ Bold, Center                   │')
    print('   │ Bảng cell           │ VAlign Center, padding         │')
    print('   │ Hình ảnh            │ Center, no indent              │')
    print('   │ Caption             │ Center, Italic, 13pt           │')
    print('   │ Header/Footer       │ Times New Roman 11pt           │')
    print('   └──────────────────────┴────────────────────────────────┘')


if __name__ == '__main__':
    format_thesis()
